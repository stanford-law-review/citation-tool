import os
import re
import warnings
import subprocess
import shutil

from bs4 import BeautifulSoup, builder
from docx import Document

# Suppress warnings related to parsing XML as HTML in BeautifulSoup
warnings.filterwarnings("ignore", category=builder.XMLParsedAsHTMLWarning)

# File paths for the footnotes and document body in the unzipped .docx structure
FTNTFILE = "temp/word/footnotes.xml"
BODYFILE = "temp/word/document.xml"


def format_url(url):
    """
    Formats a URL into an HTML anchor link.

    Parameters:
    - url (str): The URL to format.

    Returns:
    - str: The formatted URL as an HTML <a> tag.
    """
    return f'<a href="{url}">{url}</a>'


def write_italics(soup):
    """
    Adds HTML <em> tags around italicized text in the BeautifulSoup object.

    Parameters:
    - soup (BeautifulSoup): The BeautifulSoup object representing the document's XML.

    Returns:
    - BeautifulSoup: The updated soup object with <em> tags for italics.
    """
    italics = []
    formatting = soup.findAll("w:r")

    # Identify text runs containing italic formatting
    for run in formatting:
        if "w:i>" in str(run):
            italics.append(run)

    # Wrap italicized text with <em> tags
    for item in italics:
        extract1 = item.find("w:i")
        if extract1:
            extract1.extract()
        text = item.find("w:t")
        if text:
            text.contents[0].replaceWith("<em>" + str(text.contents[0]) + "</em>")

    return soup


def write_smallcaps(soup):
    """
    Adds HTML <span class="citation"> tags around small caps text in the BeautifulSoup object.

    Parameters:
    - soup (BeautifulSoup): The BeautifulSoup object representing the document's XML.

    Returns:
    - BeautifulSoup: The updated soup object with small caps wrapped in <span> tags.
    """
    smallcaps = []
    formatting = soup.findAll("w:r")

    # Identify text runs containing small caps formatting
    for run in formatting:
        if "w:smallcaps" in str(run):
            smallcaps.append(run)

    # Wrap small caps text with <span> tags
    for item in smallcaps:
        extract1 = item.find("w:smallcaps")
        if extract1:
            extract1.extract()
        text = item.find("w:t")
        if text:
            text.contents[0].replaceWith(
                '<span class="citation">' + str(text.contents[0]) + "</span>"
            )

    return soup


def write_nbsp(soup):
    """
    Replaces non-breaking spaces in the BeautifulSoup object with HTML &nbsp; entities.

    Parameters:
    - soup (BeautifulSoup): The BeautifulSoup object representing the document's XML.

    Returns:
    - BeautifulSoup: The updated soup object with &nbsp; entities.
    """
    nbsp = []
    formatting = soup.findAll("w:r")

    # Identify text runs containing non-breaking spaces
    for run in formatting:
        if "w:t" in str(run):
            nbsp.append(run)

    # Replace non-breaking spaces with &nbsp;
    for item in nbsp:
        text = item.find("w:t")
        if text and text.contents[0] == " ":
            text.contents[0].replaceWith("&nbsp;")

    return soup


def write_crossref(soup):
    """
    Removes cross-references and retains only the reference number in the BeautifulSoup object.

    Parameters:
    - soup (BeautifulSoup): The BeautifulSoup object representing the document's XML.

    Returns:
    - BeautifulSoup: The updated soup object with cross-references removed.
    """
    crossref = []
    formatting = soup.findAll("w:r")

    # Identify text runs containing cross-references
    for run in formatting:
        if "w:instrtext" in str(run):
            crossref.append(run)

    # Remove the linking code, keeping only the reference number
    for item in crossref:
        extract1 = item.find("w:instrtext")
        if extract1:
            extract1.extract()

    return soup


class Doc:
    def __init__(self, filename, num_acknowledgment_footnotes, enable_markup):
        """
        Initializes the Doc object by extracting footnotes and document text from the .docx file.

        Parameters:
        - filename (str): The path to the .docx file.
        - num_acknowledgment_footnotes (int): The number of acknowledgment footnotes to skip.
        - enable_markup (bool): Flag to enable or disable HTML markup for text formatting.
        """
        self.filename = filename
        self.enable_markup = enable_markup
        self.num_acknowledgment_footnotes = num_acknowledgment_footnotes
        self.unzip_file()  # Unzip the .docx file
        self.footnote_list = self.extract_footnotes()  # Extract footnotes
        self.doc_text = self.read_doc_text()  # Read and process document text

    def unzip_file(self):
        """
        Unzips the .docx file and extracts its contents into the 'temp' directory.
        """
        # Clean up any existing temp directory
        if os.path.exists("temp"):
            shutil.rmtree("temp")
        
        # Create the temp directory
        os.makedirs("temp", exist_ok=True)
        
        # Check if file exists and is readable
        if not os.path.exists(self.filename):
            raise FileNotFoundError(f"File not found: {self.filename}")
        
        if not os.access(self.filename, os.R_OK):
            raise PermissionError(f"Cannot read file: {self.filename}")
        
        # Check file extension
        if not self.filename.lower().endswith('.docx'):
            raise ValueError(f"File must be a .docx file, got: {self.filename}")
        
        try:
            # Use subprocess instead of os.system for better error handling
            result = subprocess.run(
                ["unzip", "-o", self.filename, "-d", "temp"],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"Failed to unzip {self.filename}. Error: {result.stderr}")
                
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"Unzip operation timed out for {self.filename}")
        except FileNotFoundError:
            raise RuntimeError("unzip command not found. Please install unzip utility.")
        except Exception as e:
            raise RuntimeError(f"Failed to unzip {self.filename}: {str(e)}")
        
        # Verify that the basic document structure exists
        if not os.path.exists("temp/word/document.xml"):
            raise RuntimeError(f"Invalid .docx structure in {self.filename}. document.xml not found.")

    def extract_footnotes(self):
        """
        Extracts footnotes from the footnotes.xml file and processes them.

        Returns:
        - list: A list of tuples, each containing the footnote number and text.
        """
        all_footnotes = []
        
        # Check if footnotes.xml exists
        if not os.path.exists(FTNTFILE):
            print(f"Warning: No footnotes found in {self.filename} (footnotes.xml not present)")
            return all_footnotes
        
        try:
            with open(FTNTFILE, "r") as f:
                soup = BeautifulSoup(f, features="lxml")
                soup = write_crossref(soup)  # Remove cross-references
                if self.enable_markup:
                    soup = write_nbsp(soup)  # Replace non-breaking spaces
                    soup = write_italics(soup)  # Format italics
                    soup = write_smallcaps(soup)  # Format small caps

                # Extract footnotes
                footnotes = soup.findAll("w:footnote")
                for fts in footnotes:
                    footnote_num = int(fts["w:id"]) - self.num_acknowledgment_footnotes
                    if footnote_num > 0:
                        footnote_text = fts.text.lstrip(
                            " .\u0002"
                        )  # Clean up footnote text
                        all_footnotes.append((str(footnote_num), footnote_text))
        
        except Exception as e:
            print(f"Warning: Error reading footnotes from {self.filename}: {e}")
            return []

        return all_footnotes

    def read_doc_text(self):
        """
        Reads and processes the main document text, applying formatting for footnotes,
        non-breaking spaces, italics, and small caps if markup is enabled.

        Returns:
        - str: The processed document text.
        """
        doc = Document(self.filename)
        if self.enable_markup:
            # Iterate over all paragraphs and apply formatting
            for para in doc.paragraphs:
                for run in para.runs:
                    # Handle footnote references
                    if (
                        run.font.superscript and run.text == ""
                    ) or run.style.name.lower() in [
                        "_noterefintext",
                        "footnote reference",
                    ]:
                        run.text = "insertfootnotehere"
                    # Handle non-breaking spaces
                    if run.text == " ":
                        run.text = "&nbsp;"
                    # Format italic text
                    if run.italic:
                        run.italic = None
                        run.text = "<em>" + run.text + "</em>"
                    # Format small caps text
                    if run.font.small_caps:
                        run.font.small_caps = None
                        run.text = '<span class="citation">' + run.text + "</span>"

        # Join paragraphs with double newlines and tab indentation
        return "\n\n\t".join([para.text for para in doc.paragraphs])

    def inject_footnote_markup(self):
        """
        Inserts formatted footnotes into the document text at the appropriate locations.

        Returns:
        - str: The document text with footnotes injected.
        """
        doc_text = self.doc_text
        doc_text = doc_text.replace(
            "insertfootnotehere", "", self.num_acknowledgment_footnotes
        )
        for footnote_num, footnote_text in self.footnote_list:
            footnote_html = f'<span class="footnote">[footnote {footnote_num}]{footnote_text}[/footnote]</span>'
            doc_text = doc_text.replace("insertfootnotehere", footnote_html, 1)
        return doc_text

    def numbered_footnotes(self):
        """
        Returns the footnotes as a list of numbered strings.

        Returns:
        - list: A list of strings where each string contains a footnote number and text.
        """
        return [" ".join(footnote) for footnote in self.footnote_list]

    def body(self):
        """
        Returns the document body text. If markup is enabled, footnotes will be injected and URLs formatted.

        Returns:
        - str: The formatted document body text.
        """
        if self.enable_markup:
            text = self.inject_footnote_markup()
            urls = re.findall(r"https:\/\/perma.cc\/\S\S\S\S-\S\S\S\S", text)
            for url in urls:
                fmt_url = format_url(url)
                text = text.replace(url, fmt_url)
            return text
        else:
            return self.doc_text

    def cleanup(self):
        """
        Cleans up the temporary directory used for extracting the .docx file.
        """
        if os.path.exists("temp"):
            shutil.rmtree("temp")
